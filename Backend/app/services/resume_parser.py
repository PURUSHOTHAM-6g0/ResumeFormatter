import PyPDF2
import httpx
from app.config import AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_KEY
import json
import re
import tempfile
import os
import base64  # For encoding images
import fitz  # PyMuPDF
from PIL import Image
import io
from docx import Document  # For DOCX text extraction
import subprocess
import platform

def extract_text_from_pdf(file_path: str) -> str:
    """Legacy function to extract text from PDF - kept for backward compatibility"""
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""

        return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files using python-docx"""
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Extract text from headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for para in section.header.paragraphs:
                    text += para.text + "\n"
            
            # Footer  
            if section.footer:
                for para in section.footer.paragraphs:
                    text += para.text + "\n"
        
        return text.strip()
        
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        raise RuntimeError(f"Failed to extract text from DOCX file: {str(e)}")

def convert_docx_to_pdf(docx_path: str) -> str:
    """Convert DOCX file to PDF using Aspose.Words first, then fallback to original methods"""
    try:
        # Create a temporary PDF file
        pdf_path = docx_path.replace('.docx', '.pdf').replace('.doc', '.pdf')
        
        # Try Aspose.Words first
        try:
            import aspose.words as aw
            
            print(f"Attempting conversion with Aspose.Words: {docx_path}")
            
            # Load the document
            doc = aw.Document(docx_path)
            
            # Save as PDF with basic settings
            doc.save(pdf_path)
            
            # Verify the PDF was created and has content
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print(f"Successfully converted DOCX to PDF using Aspose.Words: {pdf_path}")
                return pdf_path
            else:
                print("Aspose.Words created empty or invalid PDF, falling back")
                raise Exception("Invalid PDF created")
                
        except ImportError:
            print("Aspose.Words not installed, falling back to original methods")
        except Exception as e:
            print(f"Aspose.Words conversion failed: {str(e)}, falling back to original methods")
        
        # Fallback to original methods
        system = platform.system().lower()
        
        if system == "linux":
            # Use LibreOffice on Linux
            try:
                # Get the directory where the DOCX file is located
                docx_dir = os.path.dirname(docx_path)
                
                # Run LibreOffice headless conversion
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', docx_dir, docx_path
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    print(f"Successfully converted DOCX to PDF using LibreOffice")
                    return pdf_path
                else:
                    print(f"LibreOffice conversion failed: {result.stderr}")
                    raise Exception("LibreOffice conversion failed")
                    
            except (subprocess.TimeoutExpired, FileNotFoundError) as e:
                print(f"LibreOffice not available or timeout: {str(e)}")
                raise Exception("LibreOffice conversion not available")
                
        elif system == "windows":
            # Use docx2pdf on Windows
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                print(f"Successfully converted DOCX to PDF using docx2pdf")
                return pdf_path
            except ImportError:
                print("docx2pdf not available on Windows")
                raise Exception("docx2pdf not available")
                
        else:
            # Fallback: try python-docx with reportlab for cross-platform conversion
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                
                # Extract text using existing function
                text = extract_text_from_docx(docx_path)
                
                # Create PDF
                doc = SimpleDocTemplate(pdf_path, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Split text into paragraphs and add to PDF
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = Paragraph(para_text, styles['Normal'])
                        story.append(para)
                
                doc.build(story)
                print(f"Successfully converted DOCX to PDF using reportlab")
                return pdf_path
                
            except ImportError:
                print("reportlab not available for PDF generation")
                raise Exception("No PDF conversion method available")
                
    except Exception as e:
        print(f"Error converting DOCX to PDF: {str(e)}")
        raise RuntimeError(f"Failed to convert DOCX to PDF: {str(e)}")

def convert_pdf_to_images(file_path: str, dpi: int = 300) -> list:
    """Convert PDF to a list of PIL Images using PyMuPDF (fitz)"""
    try:
        # Open the PDF
        pdf_document = fitz.open(file_path)
        images = []
        
        # Get number of pages
        page_count = len(pdf_document)
        print(f"PDF has {page_count} pages - converting all pages to images")
        
        # Convert each page to an image
        for page_num in range(page_count):
            page = pdf_document.load_page(page_num)
            
            # Set the rendering matrix for higher quality
            zoom = dpi / 72  # 72 is the default DPI for PDF
            matrix = fitz.Matrix(zoom, zoom)
            
            # Render page to a pixmap (image)
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            
            # Convert pixmap to PIL Image
            img = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
            images.append(img)
            
            print(f"Converted page {page_num + 1}/{page_count} to image: {img.width}x{img.height}")
        
        # Close the document
        pdf_document.close()
        
        print(f"Successfully converted all {page_count} pages to images")
        return images
    except Exception as e:
        print(f"Error converting PDF to images with PyMuPDF: {str(e)}")
        raise

def image_to_base64(image) -> str:
    """Convert PIL Image to base64 string"""
    buffer = io.BytesIO()
    image.save(buffer, format='PNG')
    img_str = base64.b64encode(buffer.getvalue()).decode('utf-8')
    return img_str

def extract_resume_details_with_azure(text: str) -> dict:
    """Legacy function that uses text-based extraction - kept for backward compatibility"""
    url = AZURE_OPENAI_ENDPOINT
    headers = {
        "Content-Type": "application/json",
        "api-key": AZURE_OPENAI_KEY
    }

    system_prompt = (
        "You are an expert resume parser. Extract the following fields from the resume:\n"
        "- name\n"
        "- email\n"
        "- mobile\n"
        "- skills (group related skills together, and return as a list of objects with category as the key and related skills as the value. For example: [{ 'Programming Languages': ['Java', 'C++'] }, { 'Cloud': ['AWS', 'Docker'] }])\n"
        "- education (recently passed degree/institution)\n"
        "- professional_experience (CRITICAL: Summarize to fit EXACTLY 1000 characters or less. This must be concise but comprehensive, covering key achievements and technologies. Format as array: ['point1','point2',..])\n"
        "- certifications (as a list\check for certifications with images eg: microsoft certified Technology specialist)\n"
        "- experience_data (as a list of objects with each object containing the following keys: 'company', 'startDate', 'endDate', 'role', 'clientEngagement', 'program', and 'responsibilities' which is a list of bullet points describing duties)\n"
        "- summary (brief professional summary)\n"
        "\nIMPORTANT: The professional_experience field MUST be summarized to fit within 1000 characters total (including all array elements). Prioritize most important achievements and technologies.\n"
        "Return the data as valid JSON. If there is no data available for a section, try to infer it from the resume. If not possible, return 'Not available' for that section."
    )

    user_prompt = f"Resume Text:\n{text}"

    payload = {
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.2,
        "max_tokens": 6000
    }

    try:
        response = httpx.post(url, headers=headers, json=payload, timeout=50.0)
        response.raise_for_status()
        try:
            data = response.json()
            return data["choices"][0]["message"]["content"]
        except Exception as json_error:
            print("Raw response text:", response.text)  # This will show what Azure actually returned
            raise RuntimeError(f"Failed to parse JSON: {json_error}")
    except httpx.HTTPStatusError as e:
        print("Azure returned an HTTP error:", e.response.text)
        raise RuntimeError(f"Request failed with status {e.response.status_code}: {e.response.text}")

def extract_resume_details_with_azure_vision(images: list) -> dict:
    """Extract resume details using Azure OpenAI with vision capabilities"""
    url = AZURE_OPENAI_ENDPOINT
    headers = {
        "Content-Type": "application/json",
        "api-key": AZURE_OPENAI_KEY
    }

    system_prompt = (
        "You are an expert resume parser. Extract the following fields from the resume:\n"
        "- name\n"
        "- email\n"
        "- mobile\n"
        "- skills (group related skills together, and return as a list of objects with category as the key and related skills as the value. For example: [{ 'Programming Languages': ['Java', 'C++'] }, { 'Cloud': ['AWS', 'Docker'] }])\n"
        "- education (recently passed degree/institution)\n"
        "- professional_experience (CRITICAL: Summarize to fit EXACTLY 1000 characters or less. This must be concise but comprehensive, covering key achievements and technologies. Format as array: ['point1','point2',..])\n"
        "- certifications (as a list\check for certifications with images eg: microsoft certified Technology specialist)\n"
        "- experience_data (as a list of objects with each object containing the following keys: 'company', 'startDate', 'endDate', 'role', 'clientEngagement', 'program', and 'responsibilities' which is a list of bullet points describing duties)\n"
        "- summary (brief professional summary)\n"
        "\nCRITICAL INSTRUCTIONS FOR EXPERIENCE DATA EXTRACTION:\n"
        "- EXTRACT EVERY SINGLE ROW from ALL experience tables across ALL pages\n"
        "- Each table row represents a separate job/role and should be a separate object in the experience_data array\n"
        "- Do NOT skip any rows - process EVERY visible row in experience tables\n"
        "- If a company appears multiple times with different roles, create separate objects for EACH role\n"
        "- Look for tables with columns like: Role, Location, Domain, Duration, Key Projects\n"
        "- Parse ALL rows from top to bottom, including partially visible rows\n"
        "- For multi-page tables, ensure you capture continuation rows on subsequent pages\n"
        "- If you see only partial information in a row, still include it as a separate entry\n"
        "- Pay special attention to table borders and row separators to identify individual entries\n"
        "- Count the number of rows you process and ensure you capture ALL visible experience entries\n\n"
        "PROFESSIONAL EXPERIENCE SUMMARIZATION:\n"
        "- The professional_experience field MUST be summarized to fit within 1000 characters total (including all array elements)\n"
        "- Prioritize most important achievements, technologies, and impact\n"
        "- Use concise language while maintaining key information\n"
        "- Focus on quantifiable results and technical skills\n\n"
        "TABLE PARSING STRATEGY:\n"
        "1. Identify ALL tables containing work experience information\n"
        "2. Process each table row by row from top to bottom\n"
        "3. Extract data from each column for every row\n"
        "4. Create a separate experience_data object for each row\n"
        "5. Continue processing on subsequent pages if tables span multiple pages\n\n"
        "Return the data as valid JSON. If there is no data available for a section, try to infer it from the resume. If not possible, return 'Not available' for that section."
    )

    # Process ALL pages for complete coverage of 10-page resumes
    content = [{"type": "text", "text": "Parse this complete resume. This is a 10-page document. Pay special attention to extracting ALL experience data from ALL table rows across ALL pages. DO NOT miss any rows in experience tables. IMPORTANT: Summarize professional_experience to exactly 1000 characters or less:"}]
    
    # Process ALL pages (up to 10 for complete coverage)
    max_pages = min(10, len(images))
    print(f"Processing ALL {max_pages} pages out of {len(images)} total pages for comprehensive vision analysis")
    
    for i, image in enumerate(images[:max_pages]):
        base64_image = image_to_base64(image)
        content.append({
            "type": "image_url", 
            "image_url": {"url": f"data:image/png;base64,{base64_image}"}
        })
        print(f"Added page {i+1} to vision processing for complete table extraction")
        
    payload = {
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": content}
        ],
        "temperature": 0.05,  # Very low temperature for maximum consistency
        "max_tokens": 12000   # Increased token limit significantly for longer responses
    }

    try:
        response = httpx.post(url, headers=headers, json=payload, timeout=180.0)  # Increased timeout for all pages
        response.raise_for_status()
        try:
            data = response.json()
            extracted_content = data["choices"][0]["message"]["content"]
            print(f"Azure Vision API response length: {len(extracted_content)} characters")
            return extracted_content
        except Exception as json_error:
            print("Raw response text:", response.text)
            raise RuntimeError(f"Failed to parse JSON: {json_error}")
    except httpx.HTTPStatusError as e:
        print("Azure returned an HTTP error:", e.response.text)
        raise RuntimeError(f"Request failed with status {e.response.status_code}: {e.response.text}")

def clean_json_string(raw: str):
    # Remove triple backticks and language hint (```json)
    cleaned = re.sub(r"^```json\s*|\s*```$", "", raw.strip(), flags=re.MULTILINE)
    return json.loads(cleaned)

def validate_professional_experience_length(data: dict) -> dict:
    """Validate and truncate professional_experience to fit 1000 characters"""
    if 'professional_experience' in data and isinstance(data['professional_experience'], list):
        # Calculate total length
        total_length = sum(len(str(item)) for item in data['professional_experience'])
        
        if total_length > 1000:
            print(f"Professional experience too long ({total_length} chars), truncating to 1000 chars")
            
            # Truncate items to fit within 1000 characters
            truncated_items = []
            current_length = 0
            
            for item in data['professional_experience']:
                item_str = str(item)
                if current_length + len(item_str) <= 1000:
                    truncated_items.append(item)
                    current_length += len(item_str)
                else:
                    # Add partial item if there's space
                    remaining_space = 1000 - current_length
                    if remaining_space > 10:  # Only add if meaningful space left
                        truncated_item = item_str[:remaining_space-3] + "..."
                        truncated_items.append(truncated_item)
                    break
            
            data['professional_experience'] = truncated_items
            print(f"Truncated to {len(truncated_items)} items, total length: {sum(len(str(item)) for item in truncated_items)} chars")
    
    return data

# New function for processing multiple files
def process_multiple_files(file_paths: list, use_vision: bool = True) -> list:
    """Process multiple resume files and return combined results"""
    results = []
    
    for file_path in file_paths:
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            converted_pdf_path = None
            
            # Convert DOCX to PDF if using vision processing
            if file_extension in ['.doc', '.docx'] and use_vision:
                try:
                    converted_pdf_path = convert_docx_to_pdf(file_path)
                    file_extension = '.pdf'
                    processing_path = converted_pdf_path
                except Exception as e:
                    print(f"DOCX to PDF conversion failed for {file_path}, falling back to text-based processing: {str(e)}")
                    use_vision = False
                    processing_path = file_path
            else:
                processing_path = file_path
            
            # Process using vision or text-based approach
            if use_vision and file_extension == '.pdf':
                try:
                    images = convert_pdf_to_images(processing_path)
                    extracted = extract_resume_details_with_azure_vision(images)
                    parsed = clean_json_string(extracted)
                    parsed = validate_professional_experience_length(parsed)
                    processing_method = "vision"
                except Exception as e:
                    print(f"Vision processing failed for {file_path}, falling back to text-based: {str(e)}")
                    use_vision = False
            
            if not use_vision:
                if file_extension == '.pdf':
                    text = extract_text_from_pdf(file_path)
                elif file_extension in ['.doc', '.docx']:
                    text = extract_text_from_docx(file_path)
                else:
                    raise Exception(f"Unsupported file type: {file_extension}")
                
                extracted = extract_resume_details_with_azure(text)
                parsed = clean_json_string(extracted)
                parsed = validate_professional_experience_length(parsed)
                processing_method = "text"
            
            # Add filename to the result
            parsed['filename'] = os.path.basename(file_path)
            parsed['processing_method'] = processing_method
            results.append(parsed)
            
            # Clean up converted PDF if it was created
            if converted_pdf_path and os.path.exists(converted_pdf_path) and converted_pdf_path != file_path:
                os.remove(converted_pdf_path)
                
        except Exception as e:
            print(f"Error processing file {file_path}: {str(e)}")
            results.append({
                'filename': os.path.basename(file_path),
                'error': str(e),
                'processing_method': 'failed'
            })
    
    return results
















